from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from collections import OrderedDict


def hcf(*x):  # 计算最大公约数
    smaller = min(x)
    for i in reversed(range(1, smaller + 1)):
        if list(filter(lambda j: j % i != 0, x)) == []:
            return i


"""
约定条件：
    1. 所有的标题中间不能有正文内容，否则抛出解析异常
    2. 【生成的同级标题】和【正文小标题】保持纵向，同级正文保持横向
    3. 若想自适应，得加钱，无原因不解释。或者另请高明，谢谢！！！
"""


class WordOutlineGenTable():
    """
    source_docx: 源文件名
    text_length: 正文的横向长度，默认为3
    """

    def __init__(self, source_docx, text_length=3):
        self.source_docx = source_docx
        self.text_length = text_length

    def analyse_outline(self):
        source_doc = Document(self.source_docx)
        paragraph_list = source_doc.paragraphs

        wrap_dict = {}
        doc, table = self.get_word_table_instant()
        header_coo_dict = OrderedDict()
        # index_xy = [0, 0]
        normal_text_count = 0
        coordinate_list = self.get_coordinate_list()
        # 新起一行的标志：正文末-->标题
        is_normal = False
        for index,paragraph in enumerate(paragraph_list):
            header_level = paragraph.style.name
            print(header_level, paragraph.text)
            if header_level.startswith("Normal"):
                normal_text_count += 1
                if normal_text_count > self.text_length:
                    # 1. 添加一行单元格  注：table_next有点类似单元格的列表形式，访问某个单元格采用列表方式访问
                    table_next = table.add_row().cells
                    # 2. 重置normal_text_count，保证每text_length(3)个就新建一行
                    normal_text_count = 1
                    # 3. 标志为正文
                    is_normal = True
                    # 4. 合并前面的单元格
                    for header_name, coo_value in header_coo_dict.items():
                        table.cell(coo_value[0], coo_value[1]).merge(table.cell(coo_value[0] + 1, coo_value[1]))
                        coo_value[0] = coo_value[0] + 1

                else:
                    # 当正文小标题小于3时，设置内容到表格
                    coo = coordinate_list[index]
                    table.cell(coo[0], coo[1]).text = paragraph.text
                    continue
                continue
            if header_level.startswith("Heading"):
                if is_normal:
                    row_next = table.add_row().cells
                    for header_name, coo_value in header_coo_dict.items():
                        if header_name == header_level:
                            break
                        table.cell(coo_value[0], coo_value[1]).merge(table.cell(coo_value[0] + 1, coo_value[1]))
                # 生成坐标
                coo = coordinate_list[index]
                # 将标题内容写入表格
                table.cell(coo[0], coo[1]).text = paragraph.text
                # 添加标题坐标
                # index_xy[0] = coo[0]
                # index_xy[1] = coo[1]
                header_coo_dict[header_level] = coo
        return wrap_dict

    # 返回创建出的文档表格（哪个文档的哪张表）
    def get_word_table_instant(self):
        doc = Document()
        # 涉及字体样式
        doc.styles['Normal'].font.name = u"宋体"
        doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        table = doc.add_table(rows=1, cols=4 + self.text_length, style="Table Grid")
        return doc, table

    # 生成表格坐标
    def get_coordinate_list(self):
        coordinate_list = []
        for i in range(1):
            for j in range(0, 4 + self.text_length):
                coordinate_list.append([i,j])
        return coordinate_list

    class CellWrap():
        def __init__(self, content, merge_num=None, level=None):
            # 大纲解析出的单元格内容
            self.content = content
            # 大纲解析出的单元格等级，
            # 规定：1. 对标题等级来说，必须为不同等级，即1级后面必须是2级，2级后面必须是3级，否则抛出异常
            #      2. 对正文内容来说，同级段落纵向写入，段落里面的具体条例必须横向写入
            #  否则直接抛出解析异常
            self.level = level
            # 等级的数量，例如1级标题占用20个纵向单元格
            self.merge_num = merge_num
            # 上一级是谁
            self.parent_cell = None
            # 下一级节点
            self.next_cell = []
            # 同级节点的list
            self.same_level_cell = []
            # 判断是否纵向合并，一般标题默认True纵向，正文False横向
            self.portrait_flag = True


if __name__ == "__main__":
    print(hcf(10, 25, 35, 65))

    word_table = WordOutlineGenTable("./result-jiekou.docx")
    word_table.analyse_outline()
    pass
