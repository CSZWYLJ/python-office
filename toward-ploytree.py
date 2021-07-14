from docx import Document
from docx.shared import Pt  # 用来设置字体的大小
from docx.shared import Inches
from docx.oxml.ns import qn  # 设置字体
from docx.shared import RGBColor  # 设置字体的颜色
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 设置对其方式
from queue import Queue  # 用于遍历列，防止重复访问
import copy
from collections import OrderedDict
import time
from functools import wraps


class TableGenWordOutline:
    def __init__(self, source_file):
        self.source_file = source_file
        # 产生的新文件名
        self.new_file = source_file.replace("./","./result-")
        # 用于记录第一次进入同级标题的index
        self.order_list = []

    def fn_timer(function):
        @wraps(function)
        def function_timer(*args, **kwargs):
            t0 = time.time()
            result = function(*args, **kwargs)
            t1 = time.time()
            print("Total time running %s: %s seconds" % (function.__name__, str(t1 - t0)))
            return result

        return function_timer

    def analyse_table(self, table):
        column_dict = OrderedDict()
        print(f"当前表格行列信息：{len(table.rows)}行，{len(table.columns)}列")
        for index, column in enumerate(table.columns, start=1):
            column_dict[index] = self.get_merge_cell_list(column.cells[2:], index)
        return column_dict

    # 返回merge单元格列表
    def get_merge_cell_list(self, column_cells, index):
        # 用来判断是否在merge-cell是否存在
        column_cell_list = []
        # 用于返回结果
        wrap_list = []
        # 统计合并单元格的实际数量，用于后面确定父子级关系
        count = 0
        header = None
        for cell in column_cells:
            # 当这个cell表格对象不在list列表中时 TODO 加断点
            if self.not_inside_wrap_list(cell, column_cell_list):
                count = 0
                column_cell_list.append(cell)
                # 添加HeaderWrap对象的步骤
                header = HeaderWrap(cell=cell)
                wrap_list.append(header)
            count += 1
            # 设置对象的长度
            header.length = count
            if index == 1:
                header.level = 1
        return wrap_list

    @fn_timer
    def assign_value_to_wrap(self, column_dict):
        for index, wrap_list in column_dict.items():
            if index == 1:
                for first_wrap in column_dict[index]:
                    # 指定序号为1的列表元素的标题等级都为一级标题
                    first_wrap.level = 1
                continue
            if wrap_list is not None and len(wrap_list) > 0:
                wrap_list_copy = copy.copy(wrap_list)
                sub_length_sum = 0
                # TODO 循环行
                base_wrap_list = column_dict[index - 1]
                # 遍历行
                for base_wrap in base_wrap_list:
                    while True:
                        probable_sub_wrap = wrap_list_copy.pop(0)
                        # 同级情况时
                        if probable_sub_wrap.length == base_wrap.length:
                            # 若已经赋值，就跳过
                            if probable_sub_wrap.level is not None:
                                break
                            # 找到原列表为此值的引用地址
                            src_wrap = wrap_list[wrap_list.index(probable_sub_wrap)]
                            src_wrap.level = base_wrap.level
                            src_wrap.parent_cell = base_wrap.parent_cell
                            # 置为内容属性
                            base_wrap.access_flag = True
                            src_wrap.access_flag = True
                            # base_wrap.parent_cell.next_cell.append(probable_sub_wrap)
                            base_wrap.same_level_cell.append(src_wrap)
                            break
                        if sub_length_sum < base_wrap.length:
                            # 可能子节点确定是父节点的子节点时：
                            # 1，向父节点添加子节点元素  2，将子节点的parent_cell属性标志为父节点地址
                            # if len(base_wrap.next_cell) == 0:
                            base_wrap.next_cell.append(probable_sub_wrap)
                            # 找到原列表为此值的引用地址
                            src_wrap = wrap_list[wrap_list.index(probable_sub_wrap)]
                            src_wrap.parent_cell = base_wrap
                            src_wrap.level = base_wrap.level + 1
                            sub_length_sum = sub_length_sum + probable_sub_wrap.length
                            # 累加过程中如果出现了相等
                            if sub_length_sum == base_wrap.length:
                                src_wrap.parent_cell = base_wrap
                                src_wrap.level = base_wrap.level + 1
                                # base_wrap.next_cell.append(probable_sub_wrap)
                                sub_length_sum = 0
                                break
                        else:
                            raise Exception("寻找子节点的解析异常")
            else:
                raise Exception("给wrap赋值时解析异常")

    @fn_timer
    def write_doc(self, column_dict):
        doc = Document()
        for wrap in column_dict[1]:
            queue1 = Queue()
            queue2 = Queue()
            queue1.put_nowait(wrap)
            while True:
                # 当某个队列存在至少一个wrap对象时，就从该队列取值
                if queue1.qsize() > 0:
                    parent_wrap = queue1.get_nowait()
                else:
                    # print("解析完毕,queue1大小为：",queue1.qsize())
                    break
                if parent_wrap.access_flag:
                    self.set_content(parent_wrap, doc)
                else:
                    self.set_title_property(parent_wrap, doc)
                sub_wrap_list = parent_wrap.next_cell
                # TODO 为什么不能放在continue的下面？？？？？
                same_cell = parent_wrap.same_level_cell
                # 若有子集就将子集入队
                if len(sub_wrap_list) > 1:
                    for sub_wrap in sub_wrap_list:
                        queue2.put_nowait(sub_wrap)
                    while queue1.qsize() > 0:
                        queue2.put_nowait(queue1.get_nowait())
                    # 清空q2，把q2的值出队给q1，保证q1一直有值
                    while queue2.qsize() > 0:
                        queue1.put_nowait(queue2.get_nowait())
                    continue
                # 没有的set属性值,这里的same_cell_link是list，而且只有一个值，所以取第一个元素
                elif len(sub_wrap_list) == 0 and len(same_cell) == 1:
                    self.set_content(same_cell[0], doc)
                    same_cell_link = same_cell[0].same_level_cell
                    while len(same_cell_link) == 1:
                        self.set_content(same_cell_link[0], doc)
                        same_cell_link = same_cell_link[0].same_level_cell
                elif len(sub_wrap_list) == 0 and len(same_cell) == 0:
                    break
                else:
                    raise Exception("解析异常")
        doc.save("res-JKB.docx")
        print("写入完毕")

    # 判断当前的cell是否不在List中
    def not_inside_wrap_list(self, cell, topic_list):
        # 每列的第一次都为空，就表示不在里面
        if len(topic_list) == 0:
            return True
        if cell in topic_list:
            return False
        else:
            return True

    # 设置内容属性
    def set_content(self, wrap, doc):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)  # 设置段前 0 磅
        p.paragraph_format.space_after = Pt(0)  # 设置段后 0 磅
        p.paragraph_format.line_spacing = 1.5  # 设置行间距为 1.5倍
        # p.paragraph_format.first_line_indent=Inches(0.5) #段落首行缩进为 0.5英寸
        p.paragraph_format.first_line_indent = Inches(0.3346457)  # 相当于小四两个字符的缩进
        p.paragraph_format.left_indent = Inches(0)  # 设置左缩进 1英寸
        p.paragraph_format.right_indent = Inches(0)  # 设置右缩进 0.5 英寸

        text = "正文默认值" if wrap.cell.text == '' or wrap.cell.text.isspace() else wrap.cell.text
        r = p.add_run(f"{text}")
        r.font.name = u'宋体'  # 设置为宋体
        r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置为宋体，和上边的一起使用
        r.font.size = Pt(12)  # 设置字体大小为12磅 相当于 小四
        r.font.color.rgb = RGBColor(0, 0, 0)  # 设置颜色为黑色

    # 设置标题属性
    def set_title_property(self, wrap, doc):
        # doc = Document()
        if wrap.level is None:
            raise Exception("解析错误，程序未解析出标题等级")
        para_heading = doc.add_heading('', level=wrap.level)
        para_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 设置为左对齐
        para_heading.paragraph_format.space_before = Pt(0)  # 设置段前 0 磅
        para_heading.paragraph_format.space_after = Pt(0)  # 设置段后 0 磅
        para_heading.paragraph_format.line_spacing = 1.5  # 设置行间距为 1.5
        para_heading.paragraph_format.left_indent = Inches(0)  # 设置左缩进 1英寸
        para_heading.paragraph_format.right_indent = Inches(0)  # 设置右缩进 0.5 英寸

        text = "标题默认值" if wrap.cell.text == '' or wrap.cell.text.isspace() else wrap.cell.text
        run = para_heading.add_run(u"{}".format(text))
        run.font.name = u'宋体'  # 设置为宋体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置为宋体，和上边的一起使用
        run.font.size = Pt(12)  # 设置1级标题文字的大小为“小四” 为12磅
        run.font.color.rgb = RGBColor(0, 0, 0)  # 设置颜色为黑色

    @fn_timer
    def task_engine(self):
        source_doc = Document(self.source_file)
        table = source_doc.tables[0]
        result_dict = self.analyse_table(table)
        self.assign_value_to_wrap(result_dict)
        self.write_doc(result_dict)
        print("当前word表格大纲写入完成")


class HeaderWrap():
    def __init__(self, cell, length=None, level=None):
        # 某一级标题的list
        self.cell = cell
        # 该级的上一级的cell地址，里面包含了上一级cell的标题等级
        self.level = level
        # 等级的数量
        self.length = length
        # 上一级是谁
        self.parent_cell = None
        # 下一级节点
        self.next_cell = []
        # 同级节点的list
        self.same_level_cell = []
        # 判断是否为标题，若是内容：True  标题则为False
        self.access_flag = False


if __name__ == "__main__":
    # write_table = TableGenWordOutline("./jiekou.docx")
    write_table = TableGenWordOutline("./JKB.docx")
    write_table.task_engine()
